#!/usr/bin/env python3
"""Generate the implementation-backed SprintLab V3 specification and audit report."""

from __future__ import annotations

import json
import os
import re
import subprocess
from collections import Counter, defaultdict
from datetime import date
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/joshuaacha/Desktop/sprintlab")
BUILD = ROOT / "docs" / "v3" / "build"
OUTPUT = ROOT / "docs" / "v3"
DATA = json.loads((BUILD / "v3-audit-data.json").read_text())

VERSION = "3.0"
CREATED = "July 30, 2026"
BASE_COMMIT = "c00bd885e64d52e624b3cf38cc15ab3260915586"
BASE_SHORT = BASE_COMMIT[:7]
REVIEWED_VERSION = f"SprintLab 1.0.0 working tree based on {BASE_SHORT}"

INK = "0B1117"
DARK = "101922"
SURFACE = "15212B"
SLATE = "52616E"
MUTED = "6F7F8C"
LIGHT = "F2F5F7"
WHITE = "FFFFFF"
LIME = "C9FF18"
LIME_DARK = "718F00"
PALE_LIME = "EFF9D2"
RED = "C83E4D"
AMBER = "C27A18"
PALE_RED = "FBEAEC"
PALE_AMBER = "FFF3DD"
BLUE = "2B6F9C"
PALE_BLUE = "EAF4FA"


SOURCES = {
    "R1": ("Haugen et al. (2019), The Training and Development of Elite Sprint Performance", "https://sportsmedicine-open.springeropen.com/articles/10.1186/s40798-019-0221-0"),
    "R2": ("World Athletics, The 400 Metres (historical coaching review)", "https://worldathletics.org/download/downloadnsa?filename=ee60d9c7-6a54-406c-8c32-7adfd7c01fad.pdf&urlslug=the-400-metres"),
    "R3": ("Myrvang et al. (2024), Longitudinal Effects of Resisted and Assisted Sprint Training", "https://sportsmedicine-open.springeropen.com/articles/10.1186/s40798-024-00777-7"),
    "R4": ("Loturco et al. (2023), Speed Training Practices of Brazilian Olympic Sprint and Jump Coaches", "https://pmc.ncbi.nlm.nih.gov/articles/PMC10694730/"),
    "R5": ("Loturco et al. (2023), Strength and Conditioning Practices of Brazilian Olympic Sprint and Jump Coaches", "https://pmc.ncbi.nlm.nih.gov/articles/PMC10170547/"),
    "R6": ("NSCA (2009), Youth Resistance Training Position Statement", "https://www.nsca.com/globalassets/about/position-statements/position_stand_youth_resistance_training---2009.pdf"),
    "R7": ("NSCA (2016), Long-Term Athletic Development Position Statement", "https://www.nsca.com/globalassets/about/position-statements/nsca_position_statement_long-term_athletic_development.pdf"),
    "R8": ("IOC (2024), Consensus Statement on Elite Youth Athletes", "https://bjsm.bmj.com/content/58/17/946"),
    "R9": ("van Dyk et al. (2019), Nordic Hamstring Exercise and Injury Prevention", "https://pubmed.ncbi.nlm.nih.gov/30808663/"),
    "R10": ("Saw, Main, and Gastin (2016), Subjective Monitoring of Athlete Training Response", "https://bjsm.bmj.com/content/50/5/281"),
    "R11": ("Schwellnus et al. (2016), IOC Consensus on Load and Injury/Illness Risk", "https://bjsm.bmj.com/content/50/17/1043"),
    "R12": ("Jeffreys (2007), Warm-up Revisited: the RAMP Method", "https://pure.southwales.ac.uk/en/publications/warm-up-revisited-the-ramp-method-of-optimizing-performance-prepa/"),
    "R13": ("Zhang et al. (2026), Plyometric Training in Adolescent Team Sports", "https://pubmed.ncbi.nlm.nih.gov/41647156/"),
    "R14": ("Lesinski et al. (2016), Resistance Training in Youth Athletes", "https://pubmed.ncbi.nlm.nih.gov/26851290/"),
    "R15": ("Harrison et al. (2024), Muscle Damage Effects on Sprint and Change of Direction", "https://pubmed.ncbi.nlm.nih.gov/38952917/"),
    "R16": ("Prieske et al. (2023), Core Training and Performance", "https://pmc.ncbi.nlm.nih.gov/articles/PMC10588579/"),
    "R17": ("Luo et al. (2025), Core Training and Athletic Performance", "https://pmc.ncbi.nlm.nih.gov/articles/PMC12048976/"),
    "S1": ("Clark et al., NFL Combine 40-Yard Dash and Maximum Velocity", "https://pubmed.ncbi.nlm.nih.gov/28658072/"),
    "S2": ("Rumpf et al. (2016), Sprint Training Methods by Distance", "https://pubmed.ncbi.nlm.nih.gov/26492101/"),
    "S3": ("Haugen et al. (2019), Elite Sprint Performance", "https://pubmed.ncbi.nlm.nih.gov/31754845/"),
    "S7": ("Lockie et al. (2013), Acceleration Stance Kinetics", "https://pubmed.ncbi.nlm.nih.gov/23222091/"),
    "S11": ("Behm and Chaouachi (2011), Static and Dynamic Stretching", "https://pubmed.ncbi.nlm.nih.gov/21373870/"),
    "S12": ("Mayhew et al. (2010), Hand and Electronic 40-Yard Timing", "https://pubmed.ncbi.nlm.nih.gov/20072055/"),
    "S13": ("Mann et al. (2015), Reliability of 40-Yard Timing", "https://pubmed.ncbi.nlm.nih.gov/25785707/"),
    "S14": ("Haugen and Buchheit (2016), Sprint Performance Monitoring", "https://pubmed.ncbi.nlm.nih.gov/26660758/"),
    "S15": ("Abt et al. (2011), Sprint-to-Rest Ratio and Recovery Mode", "https://pubmed.ncbi.nlm.nih.gov/21386729/"),
    "S16": ("Falch et al. (2019), Change-of-Direction Training", "https://pubmed.ncbi.nlm.nih.gov/31858292/"),
}


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in edges:
            continue
        edge_data = edges.get(edge)
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key in ["val", "sz", "space", "color"]:
            if key in edge_data:
                element.set(qn("w:{}".format(key)), str(edge_data[key]))


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    node = OxmlElement("w:keepNext")
    p_pr.append(node)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    tr_pr.append(node)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_hyperlink(paragraph, text: str, url: str, color=BLUE, underline=True):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    if color:
        color_node = OxmlElement("w:color")
        color_node.set(qn("w:val"), color)
        r_pr.append(color_node)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths: list[float]):
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_repeat_headers_footers(doc: Document):
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.text = "SPRINTLAB  /  WORKOUT LIBRARY & SEASON ENGINE V3"
        p.style = doc.styles["Small Label"]
        p.runs[0].font.color.rgb = rgb(SLATE)
        p.paragraph_format.space_after = Pt(0)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = "Implementation specification  ·  "
        fp.style = doc.styles["Small Label"]
        fp.runs[0].font.color.rgb = rgb(SLATE)
        add_page_number(fp)


def setup_document(title: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.06

    for style_name, size, color, before, after in [
        ("Title", 34, INK, 0, 8),
        ("Heading 1", 20, INK, 18, 8),
        ("Heading 2", 14, INK, 13, 5),
        ("Heading 3", 11, DARK, 8, 3),
    ]:
        st = styles[style_name]
        st.font.name = "Aptos Display"
        st.font.bold = True
        st.font.size = Pt(size)
        st.font.color.rgb = rgb(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    if "Small Label" not in styles:
        s = styles.add_style("Small Label", WD_STYLE_TYPE.PARAGRAPH)
    s = styles["Small Label"]
    s.font.name = "Aptos"
    s.font.size = Pt(7.5)
    s.font.bold = True
    s.font.color.rgb = rgb(LIME_DARK)
    s.font.all_caps = True
    s.paragraph_format.space_after = Pt(3)
    s.paragraph_format.keep_with_next = True

    if "Lead" not in styles:
        s = styles.add_style("Lead", WD_STYLE_TYPE.PARAGRAPH)
    s = styles["Lead"]
    s.font.name = "Aptos"
    s.font.size = Pt(13)
    s.font.color.rgb = rgb(SLATE)
    s.paragraph_format.space_after = Pt(8)
    s.paragraph_format.line_spacing = 1.1

    if "Compact" not in styles:
        s = styles.add_style("Compact", WD_STYLE_TYPE.PARAGRAPH)
    s = styles["Compact"]
    s.font.name = "Aptos"
    s.font.size = Pt(8.3)
    s.paragraph_format.space_after = Pt(2.5)
    s.paragraph_format.line_spacing = 1.0

    if "Code Block" not in styles:
        s = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    s = styles["Code Block"]
    s.font.name = "Menlo"
    s.font.size = Pt(7.8)
    s.font.color.rgb = rgb(LIGHT)
    s.paragraph_format.left_indent = Inches(0.14)
    s.paragraph_format.right_indent = Inches(0.14)
    s.paragraph_format.space_before = Pt(4)
    s.paragraph_format.space_after = Pt(6)
    s.paragraph_format.line_spacing = 1.0

    doc.core_properties.title = title
    doc.core_properties.subject = "SprintLab implementation specification"
    doc.core_properties.author = "SprintLab / Joshua Acha"
    doc.core_properties.keywords = "SprintLab, workout library, season engine, planner, implementation"
    set_repeat_headers_footers(doc)
    return doc


def add_cover(doc: Document, subtitle: str):
    p = doc.add_paragraph(style="Small Label")
    p.add_run("SCIENCE-BASED SPEED TRAINING  /  IMPLEMENTATION SPECIFICATION")
    p.paragraph_format.space_before = Pt(45)

    title = doc.add_paragraph(style="Title")
    title.add_run("SprintLab Workout\nLibrary and Season Engine")
    v = doc.add_paragraph()
    v.paragraph_format.space_after = Pt(18)
    r = v.add_run("V3")
    r.font.name = "Aptos Display"
    r.font.size = Pt(52)
    r.font.bold = True
    r.font.color.rgb = rgb(LIME_DARK)

    lead = doc.add_paragraph(style="Lead")
    lead.add_run(subtitle)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.8)
    shade(cell, DARK)
    set_cell_margins(cell, top=180, start=200, bottom=180, end=200)
    set_cell_border(cell, top={"val": "single", "sz": 18, "color": LIME}, left={"val": "single", "sz": 18, "color": LIME}, bottom={"val": "single", "sz": 18, "color": LIME}, right={"val": "single", "sz": 18, "color": LIME})
    for label, value in [
        ("VERSION", VERSION),
        ("CREATED", CREATED),
        ("IMPLEMENTATION REVIEWED", REVIEWED_VERSION),
        ("BASE COMMIT", BASE_COMMIT),
        ("SOURCE OF TRUTH", "Current repository behavior and passing planner regressions"),
    ]:
        p = cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
        a = p.add_run(f"{label}\n")
        a.font.size = Pt(7.5)
        a.font.bold = True
        a.font.color.rgb = rgb(LIME)
        b = p.add_run(value)
        b.font.size = Pt(10)
        b.font.color.rgb = rgb(WHITE)
        p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.add_run("Prepared from the current SprintLab working tree. V2 is background material only.").italic = True
    p.runs[0].font.color.rgb = rgb(SLATE)
    doc.add_page_break()


def add_heading(doc: Document, text: str, level=1, label: str | None = None):
    if label:
        doc.add_paragraph(label, style="Small Label")
    return doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: Iterable[str], level=0, compact=False):
    for item in items:
        p = doc.add_paragraph(style="Compact" if compact else "Normal")
        p.style = doc.styles["Compact"] if compact else doc.styles["Normal"]
        p.paragraph_format.left_indent = Inches(0.22 + 0.18 * level)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        p.add_run("• ").font.color.rgb = rgb(LIME_DARK)
        p.add_run(item)


def add_numbered(doc: Document, items: Iterable[str]):
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.22)
        r = p.add_run(f"{index:02d}  ")
        r.font.bold = True
        r.font.color.rgb = rgb(LIME_DARK)
        p.add_run(item)


def add_callout(doc: Document, title: str, text: str, tone="info"):
    colors = {
        "info": (PALE_BLUE, BLUE),
        "success": (PALE_LIME, LIME_DARK),
        "warning": (PALE_AMBER, AMBER),
        "error": (PALE_RED, RED),
    }
    fill, border = colors[tone]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.85)
    shade(cell, fill)
    set_cell_margins(cell, top=120, start=150, bottom=120, end=150)
    set_cell_border(cell, left={"val": "single", "sz": 24, "color": border})
    p = cell.paragraphs[0]
    r = p.add_run(title.upper())
    r.font.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = rgb(border)
    p.paragraph_format.space_after = Pt(3)
    p2 = cell.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(0)
    return table


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], font_size=7.6):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    prevent_row_split(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        shade(cell, DARK)
        set_cell_margins(cell, 85, 90, 85, 90)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        r.font.bold = True
        r.font.color.rgb = rgb(WHITE)
        r.font.size = Pt(font_size)
    for row_index, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        if row_index % 2:
            for cell in row.cells:
                shade(cell, "F5F7F8")
        for i, text in enumerate(values):
            cell = row.cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell, 75, 85, 75, 85)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(text))
            r.font.size = Pt(font_size)
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def code_block(doc: Document, text: str):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, DARK)
    set_cell_margins(cell, top=120, start=130, bottom=120, end=130)
    p = cell.paragraphs[0]
    p.style = doc.styles["Code Block"]
    p.add_run(text)
    return table


def clean(value: Any) -> str:
    if value is None:
        return "—"
    if isinstance(value, list):
        return ", ".join(str(x) for x in value) if value else "—"
    return str(value)


def dose(item: dict[str, Any]) -> str:
    bits = []
    if item.get("sets") is not None:
        bits.append(f"{item['sets']} sets")
    if item.get("reps") is not None:
        bits.append(f"{item['reps']} reps")
    if item.get("distanceMeters") is not None:
        bits.append(f"{item['distanceMeters']} m")
    if item.get("fastZoneMeters") is not None:
        bits.append(f"{item['fastZoneMeters']} m fast zone")
    if item.get("durationSeconds") is not None:
        sec = item["durationSeconds"]
        bits.append(f"{sec // 60:g} min" if sec % 60 == 0 else f"{sec} s")
    if not bits:
        bits.append("completion item")
    rx = " × ".join(bits[:2]) + ((" · " + " · ".join(bits[2:])) if len(bits) > 2 else "")
    if item.get("intensity"):
        rx += f" · {item['intensity'].get('description', '')}"
    if item.get("recovery"):
        rx += f" · Rest: {item['recovery'].get('description', '')}"
    if item.get("notes"):
        rx += f" · {item['notes']}"
    return re.sub(r"\s+", " ", rx).strip(" ·")


def compact_sections(workout: dict[str, Any]) -> list[tuple[str, list[str]]]:
    result = []
    order = ["warmup", "sprintWork", "plyometrics", "strength", "coreBodyweight", "cooldown"]
    for key in order:
        section = workout["sections"][key]
        lines = [f"{item['name']}: {dose(item)}" for item in section["items"]]
        if lines:
            result.append((section["label"], lines))
    return result


def pathway_group(workout: dict[str, Any]) -> str:
    wid = workout["id"]
    if wid.startswith("F40-"):
        return "Football / 40-Yard"
    if wid.startswith("GEN-"):
        return "General Speed"
    paths = set(workout["eventPathways"])
    if paths == {"short-sprint-100-200"}:
        return "60m / 100m"
    if paths == {"long-sprint-200-400"}:
        return "200m / 400m"
    return "Shared / Support"


def scenario_selected_ids(s: dict[str, Any]) -> list[str]:
    result = s["result"]
    if result["status"] != "ready":
        return []
    return [
        item["workoutId"] + (("+" + "+".join(item["supportWorkoutIds"])) if item["supportWorkoutIds"] else "")
        for item in result["suggestions"]
    ]


def scenario_rating(s: dict[str, Any]) -> tuple[str, str]:
    name = s["scenario"]["name"]
    result = s["result"]
    if result["status"] != "ready":
        return "FAIL", "Planner did not return a complete week."
    if name.startswith("Reduced-readiness advanced 60m/100m"):
        return "FAIL", "Track pathway returned the same advanced week; weekly tier regression is not applied."
    if name.startswith("Reduced-readiness advanced 200m/400m"):
        return "FAIL", "Track pathway returned the same advanced week; weekly tier regression is not applied."
    if name.startswith("Advanced 200m/400m"):
        return "QUESTIONABLE", "Structure is coherent, but selected IDs are identical to trained and no distinct advanced long-sprint tier is expressed."
    if name.startswith("Advanced 60m/100m"):
        return "QUESTIONABLE", "Structure is coherent, but it mixes a general-speed acceleration record with shared track records."
    if name.startswith("Competition 60m/100m"):
        return "QUESTIONABLE", "Three-role competition week is complete, but a court/limited-space acceleration record is selected for a track profile."
    if name.startswith("In-season"):
        return "QUESTIONABLE", "Practice/game blocking works, but the remaining three sessions contain no meaningful loaded-strength exposure."
    return "PASS", "The generated week contains the expected pathway qualities, support structure, and explained open days."


EXPECTED = {
    ("60m/100m", "beginner", 3): "Acceleration; upright speed; one low/support day; explosive work; two foundational strength pairings.",
    ("60m/100m", "beginner", 5): "Acceleration; upright speed; two low/support days; elastic power; meaningful strength.",
    ("60m/100m", "intermediate", 3): "Acceleration; low support; maximum velocity; paired strength.",
    ("60m/100m", "intermediate", 5): "Three high-output roles, two lower-output roles, and two strength pairings.",
    ("60m/100m", "advanced", 5): "Advanced short-sprint week with acceleration, max velocity, speed endurance, support, and strength.",
    ("200m/400m", "beginner", 3): "Acceleration; introductory upright speed; support; explosive work; foundational strength.",
    ("200m/400m", "beginner", 5): "Acceleration; upright speed; two support days; elastic power; meaningful strength.",
    ("200m/400m", "intermediate", 3): "Acceleration; tempo support; maximum velocity; paired strength.",
    ("200m/400m", "intermediate", 5): "Acceleration; max velocity; controlled speed endurance; two support roles; strength.",
    ("200m/400m", "advanced", 5): "Advanced long-sprint week with distinct progression and event-specific loading.",
    ("Football/40-Yard", "beginner", 3): "Simple starts/acceleration; upright speed; support; explosive and foundational strength.",
    ("Football/40-Yard", "beginner", 5): "Two quality speed exposures; two support roles; elastic work; two strength pairings.",
    ("Football/40-Yard", "intermediate", 3): "Trained 40-yard acceleration, support, upright speed, and strength.",
    ("Football/40-Yard", "intermediate", 5): "Three quality 40-yard roles, two support roles, and two strength pairings.",
    ("Football/40-Yard", "advanced", 5): "Advanced acceleration/upright tiers, transfer exposure, support, and strength.",
    ("General Speed", "beginner", 3): "Simple acceleration, upright speed, mechanics/support, explosive work, and strength.",
    ("General Speed", "beginner", 5): "Acceleration, upright speed, two support roles, elastic work, and strength.",
    ("General Speed", "intermediate", 3): "Trained acceleration, support, upright speed, and strength.",
    ("General Speed", "intermediate", 5): "Acceleration, upright speed, integration, two support roles, and strength.",
    ("General Speed", "advanced", 5): "Advanced acceleration/upright tiers, integration, support, and strength.",
}


def expected_for(s: dict[str, Any]) -> str:
    sc = s["scenario"]
    if sc["readiness"] == "reduced":
        return "One reviewed tier lower than the otherwise supported advanced week, with the reason displayed."
    if sc["context"] == "competition":
        return "Competition-appropriate microdoses/support around competition load; no blind stacking."
    return EXPECTED.get((sc["pathway"], sc["level"], len(sc["days"])), "Coherent pathway-specific week matching the supplied schedule.")


def actual_day_summary(s: dict[str, Any]) -> list[str]:
    result = s["result"]
    if result["status"] != "ready":
        return [f"{result['status']}: {result.get('message', '')}"]
    lines = []
    for suggestion in result["suggestions"]:
        planned = suggestion["plannedWorkout"]
        track = []
        ply = []
        strength = []
        support = []
        for section in planned["sections"]:
            names = [e["name"] for e in section["exercises"]]
            if section["title"] == "Track":
                track += names
            elif "Plyometrics" in section["title"]:
                ply += names
            elif "Strength" in section["title"]:
                strength += names
            elif section["title"] in ("Warm-up", "Cooldown", "Core / bodyweight"):
                support += names
        bits = [f"{suggestion['dayIndex']} · {planned['id']} · {planned['title']}"]
        if track:
            bits.append("Sprint: " + "; ".join(track))
        if ply:
            bits.append("Explosive: " + "; ".join(ply))
        if strength:
            bits.append("Strength: " + "; ".join(strength))
        if support:
            bits.append("Support: " + "; ".join(support))
        lines.append(" | ".join(bits))
    return lines


def add_toc(doc: Document):
    add_heading(doc, "Contents", 1, "DOCUMENT MAP")
    rows = [
        ("01", "Authority, purpose, and scope"),
        ("02", "Change summary from V2"),
        ("03", "Athlete profile inputs actually used"),
        ("04", "Supported pathways"),
        ("05", "Training levels, tiers, and readiness"),
        ("06", "Current implemented schemas"),
        ("07", "Deterministic weekly-plan generation"),
        ("08", "Execution, tracking, completion, and logging"),
        ("09", "Exercise resolution and alternatives"),
        ("10", "Library approval validation"),
        ("11", "Approved workout-record inventory"),
        ("12", "Verified profile behavior"),
        ("13", "Research rationale and citations"),
        ("14", "Known limitations and planned behavior"),
        ("A", "Nonselectable draft and archived records"),
    ]
    add_table(doc, ["#", "Section"], [[a, b] for a, b in rows], [0.55, 6.2], 9)
    doc.add_page_break()


def add_schema_section(doc: Document):
    add_heading(doc, "Current implemented schemas", 1, "06  /  DATA CONTRACTS")
    doc.add_paragraph(
        "The excerpts below are simplified only for presentation. Field names and unions match the current TypeScript implementation. "
        "Optional sport-aware metadata is additive; it does not replace the core track-library fields."
    )
    add_heading(doc, "Library workout and item", 2)
    code_block(doc, """type LibraryWorkoutItem = {
  id: string; name: string;
  sets?: number; reps?: number;
  distanceMeters?: number; fastZoneMeters?: number;
  durationSeconds?: number;
  intensity?: { basis; min?; max?; description };
  recovery?: { afterRepSeconds?; afterSetSeconds?;
                rangeSeconds?; description };
  coachingCues: string[]; notes?: string;
  countsTowardSprintVolume: boolean;
  countsTowardHighIntensityVolume: boolean;
};

type LibraryWorkout = {
  id; slug; name; purpose; intendedAthlete;
  primaryCategory; secondaryCategories[];
  eventPathways[]; eventTags[]; athleteLevels[]; seasonPhases[];
  specialistProfiles[]; equipmentRequired[]; equipmentOptional[];
  surface: { required[]; preferred[]; prohibited[]; notes };
  sections: { warmup; sprintWork; plyometrics;
              strength; coreBodyweight; cooldown };
  intensitySummary; recoverySummary; metrics;
  coachingCues[]; modifications[]; safetyNotes[]; sourceNotes[];
  familyId; progressionLevel; regressionWorkoutId?;
  progressionWorkoutId?; version; approvalStatus;
  createdAt; updatedAt; archivedAt?;
  sports?; speedGoals?; speedPathways?; trainingContexts?;
  distanceUnit?; testType?; directionPattern?;
};""")
    add_heading(doc, "Planner output", 2)
    code_block(doc, """type SuggestedPlanDay = {
  dayIndex; weeklyRole; loadClass; targetCategory;
  workoutId; supportWorkoutIds[];
  plannedWorkout;
  whyThisFits[]; harderOptionsExcluded[];
  requiredSetup; stopRule; alternatives[];
};

type WeeklyPlanSuggestion =
  | { status: "ready"; schedule[]; suggestions[];
      summary; warnings[] }
  | { status: "coach-managed" | "unsupported-sport" | "no-match";
      title; message; reasons[] };""")
    add_heading(doc, "Execution and saved results", 2)
    code_block(doc, """type ExerciseTracking =
  | { kind: "completion" }
  | { kind: "track"; reps; distanceMeters?;
      targetIntensity?; restSeconds? }
  | { kind: "strength"; sets; targetReps;
      targetLoad?; restSeconds? };

type ActualExerciseResult = {
  exerciseId; sectionTitle; trackingKind; origin?;
  exerciseSnapshot?; status; changeReason?; notes;
  trackReps?: TrackRepResult[];
  strengthSets?: StrengthSetResult[];
  quickCompletionSnapshot?: {
    status; trackRepStatuses?; strengthSetStatuses?;
  };
};""")
    add_callout(
        doc,
        "Execution-structure conflict",
        "The library stores sets and reps separately. Strength execution preserves both. Sprint execution currently converts "
        "(sets × reps) into one flat rep count, so 2×5 becomes 10 sequential reps in the active session and stored result. "
        "Per-side structure is not a first-class field. V3 records this as an unresolved implementation limitation rather than claiming full fidelity.",
        "error",
    )


def add_inventory(doc: Document):
    approved = [w for w in DATA["inventory"] if w["recommendationEligible"]]
    grouped: dict[str, dict[str, list[dict[str, Any]]]] = defaultdict(lambda: defaultdict(list))
    for workout in approved:
        grouped[pathway_group(workout)][workout["primaryCategory"]].append(workout)

    add_heading(doc, "Approved workout-record inventory", 1, "11  /  SELECTABLE LIBRARY")
    doc.add_paragraph(
        f"The current seed contains {DATA['librarySummary']['total']} records: "
        f"{DATA['librarySummary']['approved']} Approved and recommendation-eligible, "
        f"{DATA['librarySummary']['draft']} Draft, and {DATA['librarySummary']['archived']} Archived. "
        "Every Approved record below passed the current implemented approval validator."
    )
    add_callout(
        doc,
        "Inventory reading rule",
        "A record can support more than one event, level, or phase. Grouping below is for navigation only; the exact eventPathways, "
        "eventTags, athleteLevels, and seasonPhases shown on each record control eligibility.",
        "info",
    )

    section_order = ["60m / 100m", "200m / 400m", "Football / 40-Yard", "General Speed", "Shared / Support"]
    record_number = 0
    for group in section_order:
        if group not in grouped:
            continue
        if record_number:
            doc.add_page_break()
        add_heading(doc, group, 2)
        categories = grouped[group]
        for category in sorted(categories):
            add_heading(doc, category.replace("-", " ").title(), 3)
            for workout in sorted(categories[category], key=lambda x: (x["progressionLevel"], x["id"])):
                record_number += 1
                table = doc.add_table(rows=1, cols=1)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = table.cell(0, 0)
                shade(cell, "F5F7F8")
                set_cell_margins(cell, 110, 135, 110, 135)
                set_cell_border(
                    cell,
                    top={"val": "single", "sz": 10, "color": LIME_DARK},
                    left={"val": "single", "sz": 10, "color": "C7D0D6"},
                    bottom={"val": "single", "sz": 10, "color": "C7D0D6"},
                    right={"val": "single", "sz": 10, "color": "C7D0D6"},
                )
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                rid = p.add_run(f"{workout['id']}  ")
                rid.font.bold = True
                rid.font.color.rgb = rgb(LIME_DARK)
                name = p.add_run(workout["name"])
                name.font.bold = True
                name.font.size = Pt(11)
                meta = cell.add_paragraph(style="Compact")
                meta.add_run(
                    f"Tier: {clean(workout['athleteLevels'])}  ·  "
                    f"Events: {clean(workout['eventTags'])}  ·  "
                    f"Phases: {clean(workout['seasonPhases'])}\n"
                    f"Pathway keys: {clean(workout['eventPathways'])}  ·  "
                    f"Duration: {workout['metrics']['estimatedDurationMinutes'][0]}–{workout['metrics']['estimatedDurationMinutes'][1]} min  ·  "
                    f"Volume: {workout['metrics']['totalSprintVolumeMeters']} m total / "
                    f"{workout['metrics']['highIntensitySprintVolumeMeters']} m high intensity"
                )
                purpose = cell.add_paragraph(style="Compact")
                r = purpose.add_run("Purpose  ")
                r.font.bold = True
                purpose.add_run(workout["purpose"])
                for label, lines in compact_sections(workout):
                    p = cell.add_paragraph(style="Compact")
                    p.paragraph_format.left_indent = Inches(0.16)
                    p.paragraph_format.first_line_indent = Inches(-0.16)
                    tag = p.add_run(f"{label}  ")
                    tag.font.bold = True
                    tag.font.color.rgb = rgb(SLATE)
                    p.add_run("  |  ".join(lines))
                p = cell.add_paragraph(style="Compact")
                p.add_run("Intensity  ").bold = True
                p.add_run(workout["intensitySummary"])
                p.add_run("   Rest  ").bold = True
                p.add_run(workout["recoverySummary"])
                p = cell.add_paragraph(style="Compact")
                p.add_run("Progression  ").bold = True
                p.add_run(
                    f"family {workout['familyId']} · level {workout['progressionLevel']} · "
                    f"regression {workout.get('regressionWorkoutId') or '—'} · progression {workout.get('progressionWorkoutId') or '—'}"
                )
                p.add_run("   Sources  ").bold = True
                p.add_run(", ".join(s["sourceId"] for s in workout["sourceNotes"]))
                if record_number % 2 == 0:
                    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_regression_matrix(doc: Document):
    add_heading(doc, "Verified profile behavior", 1, "12  /  REGRESSION MATRIX")
    doc.add_paragraph(
        "The following results were generated from the real selector on July 30, 2026. "
        "Normal profiles were healthy, general-preparation/offseason, normally recovered, had no competition that week, "
        "and used exact Monday/Wednesday/Friday or Monday/Tuesday/Wednesday/Friday/Saturday availability. "
        "Reduced-readiness profiles set currentPain=true."
    )
    for pathway in ["60m/100m", "200m/400m", "Football/40-Yard", "General Speed"]:
        add_heading(doc, pathway, 2)
        scenarios = [s for s in DATA["scenarios"] if s["scenario"]["pathway"] == pathway]
        rows = []
        for s in scenarios:
            rating, reason = scenario_rating(s)
            sc = s["scenario"]
            inputs = (
                f"{sc['level']}; {len(sc['days'])} available days; "
                f"{sc['context']}; readiness {sc['readiness']}"
            )
            rows.append([
                inputs,
                expected_for(s),
                "\n".join(scenario_selected_ids(s)) or s["result"]["status"],
                rating,
                reason,
            ])
        add_table(doc, ["Inputs", "Expected structure", "Selected IDs", "Result", "Reason"], rows, [1.25, 1.75, 1.8, 0.65, 1.45], 6.4)

    add_heading(doc, "Competition and schedule-load cases", 2)
    cases = [s for s in DATA["scenarios"] if s["scenario"]["context"] == "competition"]
    rows = []
    for s in cases:
        rating, reason = scenario_rating(s)
        rows.append([
            s["scenario"]["name"],
            expected_for(s),
            ", ".join(scenario_selected_ids(s)),
            rating,
            reason,
        ])
    add_table(doc, ["Profile", "Expected", "Selected IDs", "Result", "Reason"], rows, [1.35, 1.55, 1.9, 0.65, 1.45], 6.4)

    doc.add_page_break()
    add_heading(doc, "Complete generated outputs", 2)
    doc.add_paragraph(
        "These compact traces preserve the complete selected week: day index, planned-workout ID/title, sprint work, explosive work, "
        "strength, and support components. Open days are represented in the seven-day schedule as explained rest/open days."
    )
    for s in DATA["scenarios"]:
        p = doc.add_paragraph(style="Small Label")
        p.add_run(s["scenario"]["name"])
        for line in actual_day_summary(s):
            p = doc.add_paragraph(style="Compact")
            p.paragraph_format.left_indent = Inches(0.16)
            p.paragraph_format.first_line_indent = Inches(-0.12)
            p.add_run("• ").font.color.rgb = rgb(LIME_DARK)
            p.add_run(line)


def add_sources(doc: Document):
    add_heading(doc, "Research rationale and citations", 1, "13  /  SCIENTIFIC BASIS")
    doc.add_paragraph(
        "Implementation details are governed by the repository. Scientific rationale is drawn from saved SprintLab research documents. "
        "A source ID on a workout means the record was authored with that publication/resource in view; it does not imply endorsement of SprintLab."
    )
    used = Counter()
    for workout in DATA["inventory"]:
        for note in workout.get("sourceNotes", []):
            used[note["sourceId"]] += 1
    for source_id in sorted(used, key=lambda x: (x[0], int(re.search(r"\d+", x).group()) if re.search(r"\d+", x) else 999)):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.28)
        r = p.add_run(f"{source_id}  ")
        r.font.bold = True
        r.font.color.rgb = rgb(LIME_DARK)
        if source_id in SOURCES:
            title, url = SOURCES[source_id]
            p.add_run(f"{title}. Used by {used[source_id]} current record(s). ")
            add_hyperlink(p, url, url)
        else:
            p.add_run(
                f"Internal authoring note used by {used[source_id]} record(s). "
                "No external publication is claimed."
            )


def build_v3() -> Path:
    doc = setup_document("SprintLab Workout Library and Season Engine V3")
    add_cover(
        doc,
        "The implementation specification for SprintLab’s deterministic workout library, season logic, weekly planner, execution model, and regression behavior.",
    )
    add_toc(doc)

    add_heading(doc, "Authority, purpose, and scope", 1, "01  /  SOURCE OF TRUTH")
    doc.add_paragraph(
        "V3 replaces V2 as the implementation specification. It describes what the current code stores, selects, displays, tracks, "
        "and verifies. V2 and the research documents remain background and scientific rationale; they do not override tested behavior."
    )
    add_callout(
        doc,
        "Reviewed state",
        f"App version 1.0.0; working tree based on commit {BASE_COMMIT}. The repository had uncommitted planner, library, execution, "
        "and onboarding changes during review, so this document identifies the exact base commit and creation date rather than pretending "
        "the reviewed state is a clean commit.",
        "warning",
    )
    add_heading(doc, "What the system does", 2)
    add_bullets(doc, [
        "Stores reviewed workout records with eligibility metadata, prescriptions, six authored sections, metrics, progression links, and approval state.",
        "Derives a season phase, allocates a weekly high/low architecture, filters eligible records, ranks deterministic candidates, pairs strength, and returns a seven-day preview.",
        "Converts selected library records into executable PlannedWorkout snapshots, records per-rep or per-set results where supported, and saves completed-session logs.",
        "Applies daily readiness guidance before launch and applies one-tier weekly regression for Football/40-Yard and General Speed when current pain, return-to-training, or taper context is present.",
    ])
    add_heading(doc, "Deterministic versus personalized", 2)
    add_table(doc, ["Deterministic", "Personalized by saved inputs"], [
        ["Weekly role order; hard gates; candidate scoring; tie-breaks; approved-record selection; paired strength IDs; schedule construction.", "Event/pathway, experience tier, season phase, availability, practice/competition blocks, goal-score bonus, session-duration bonus, plan mode, and selected non-track tier regression."],
        ["Same inputs and same library state return the same week.", "The system does not use a live AI model and does not generate new workouts."],
    ], [3.35, 3.35], 8)
    add_heading(doc, "Outside current implementation", 2)
    add_bullets(doc, [
        "No live AI coaching, cloud account, Supabase backend, payments, or native-only planning dependency.",
        "No dedicated soccer, basketball, baseball/softball, lacrosse, rugby, or volleyball progression; those profiles route through the general linear-speed planner unless Football is primary.",
        "No verified performance projection or guarantee.",
        "No complete first-class representation of sprint set boundaries or per-side dosing in execution.",
    ])

    add_heading(doc, "Change summary from V2", 1, "02  /  WHAT CHANGED")
    changes = [
        ["Authority", "V2 described an intended architecture.", "V3 documents the current TypeScript implementation and passing regressions."],
        ["Library size", "Background inventory and protocols.", "70 seeded records: 64 eligible Approved, 4 Draft, 2 Archived."],
        ["Pathways", "Broad track/library concepts.", "Four documented planning pathways: 60m/100m, 200m/400m, Football/40-Yard, General Speed; shared/support records are selectable where eligible."],
        ["Week generation", "Conceptual season templates.", "Real high/low weekly architecture plus deterministic record filtering, ranking, strength pairing, and seven-day schedule output."],
        ["Access", "Access questions treated as planner constraints.", "Ordinary missing access is treated as unknown/standard; only specialized assisted/force-plate technology is gated."],
        ["Readiness", "General readiness concept.", "Daily green/yellow/red launch evaluation is separate from weekly tier regression; tier regression is currently non-track only."],
        ["Execution", "Schema intent.", "Strength sets persist; track set boundaries are currently flattened and documented as a defect."],
        ["Validation", "Broader intended approval rules.", "V3 lists only implemented checks and marks stricter dose/reference checks as planned."],
    ]
    add_table(doc, ["Area", "V2/background", "Verified V3 implementation"], changes, [1.05, 2.25, 3.45], 7.1)

    add_heading(doc, "Athlete profile inputs actually used", 1, "03  /  PLANNER INPUTS")
    add_table(doc, ["Input", "How current code uses it", "Not used / caveat"], [
        ["primarySport / sport", "Track uses the track planner; Football uses the 40-yard planner; every other non-track sport uses General Speed.", "sports[] and sportPosition are not used to specialize weekly selection."],
        ["primaryEvent", "60m/100m → short pathway; 200m/400m → long pathway; unsupported track events fall back to 100m.", "Secondary events and personal bests do not choose records."],
        ["experienceLevel", "beginner→foundation; developing→developing; intermediate→trained; advanced/elite→advanced.", "It gates eligible records and changes beginner architecture; it does not create a separate strength system."],
        ["trainingDaysPerWeek / availableTrainingDays", "Exact available days are used when present; otherwise 1–5 day defaults are filled.", "Planner caps generated roles at five."],
        ["preferredRestDay", "Blocked when preferredRestDayAnswered is not false.", "A migration default can still block unless explicitly marked unanswered."],
        ["sportPracticeDays / otherSportDays", "Blocked before session matching.", "Practices are not inserted as executable sessions; they appear as open/existing training days."],
        ["gameOrCompetitionDays / priorityMeets", "Current-week dates or recurring weekdays are blocked; meet windows filter candidates.", "gameScheduleVaries and currentTeamTrainingLoad do not directly score candidates."],
        ["busySchoolDays", "Avoided when auto-placing days, then used only if needed to fill the requested count.", "Does not lower training quality."],
        ["seasonCalendar / seasonPhaseOverride / trainingContext", "Derives general prep, specific prep, precompetition, competition, taper, or transition. Override is honored.", "Without enough calendar context the planner may return needs-calendar/no-match."],
        ["currentPain / return-to-training", "Regresses Football/General Speed one tier; also informs explanations.", "Does not currently regress track weekly tier."],
        ["usualSessionDurationMinutes", "Adds a small ranking bonus when candidate minimum duration fits.", "Does not hard-exclude longer sessions."],
        ["speedGoals", "Adds a small candidate score when record speedGoals overlap.", "Does not create a workout."],
        ["trainingPlanMode / loggingOnlyMode", "Coach/log modes suppress generated weeks.", "Library and logging remain usable."],
        ["ordinary access/equipment fields", "Displayed in setup and workout alternatives; current logistics matching assumes standard authored sessions.", "Track, gym, blocks, sled, cones, and normal weight-room answers do not gate recommendations."],
    ], [1.55, 3.35, 1.85], 6.7)

    add_heading(doc, "Supported pathways", 1, "04  /  PROGRAM SCOPE")
    add_table(doc, ["Pathway", "Current intent", "Verified selection behavior"], [
        ["60m / 100m", "Starts, acceleration, upright/max velocity, limited speed endurance, paired strength, and lower-output support.", "Track planner; EventPathway key is short-sprint-100-200. 60m and 100m share the same architecture but event tags filter records."],
        ["200m / 400m", "Acceleration, upright speed, tempo, controlled speed endurance, later special endurance, and paired strength.", "Track planner; EventPathway key is long-sprint-200-400. 200m maps here, but architecture’s long-sprint flag currently checks only 400m."],
        ["Football / 40-Yard", "Distinct tiered early acceleration and upright-speed records, then trained/advanced transfer work.", "Non-track planner using F40-ACC-01/02/03 and F40-MAX-01/02/03 plus F40-TRANSFER-01."],
        ["General Speed", "Linear acceleration, upright speed, low technical/support work, tiered records, and trained/advanced integration.", "Non-track planner using GEN-ACC-01/02/03, GEN-MAX-01/02/03, GEN-LOW-01, GEN-INTEGRATE-01 and support records."],
        ["Shared / Support", "Tempo, strength, plyometrics, core, testing, meet preparation, and records eligible across pathways.", "Chosen through shared pathway metadata, preferred IDs, category alternatives, and deterministic ranking."],
    ], [1.25, 2.65, 2.85], 7.2)

    add_heading(doc, "Training levels, tiers, and readiness", 1, "05  /  DOSAGE & REGRESSION")
    add_table(doc, ["Profile answer", "Library level", "Planner meaning"], [
        ["Beginner", "foundation", "Beginner architecture; lower-complexity fifth role; still receives loaded strength and explosive work."],
        ["Developing", "developing", "Eligible developing records; non-track tier currently resolves to foundation."],
        ["Intermediate", "trained", "Full trained weekly structure and trained tiered Football/General records."],
        ["Advanced / Elite", "advanced", "Advanced eligible records; Football/General use tier 03."],
    ], [1.45, 1.25, 4.05], 7.6)
    add_heading(doc, "Weekly tier regression", 2)
    doc.add_paragraph(
        "For Football/40-Yard and General Speed only, currentPain=true, trainingContext=return-to-training, or taper reduces an "
        "advanced tier to trained, or trained to foundation. The explanation includes: “This session was reduced one reviewed level…” "
        "The selector never chooses a record above the resulting eligible level."
    )
    add_heading(doc, "Daily readiness launch evaluation", 2)
    add_bullets(doc, [
        "Red: severe/acute pain, expected hesitation at maximum speed, or a worse warm-up reassessment. Maximal sprinting must not begin.",
        "Yellow: sleep below personal/referenced baseline, low sleep quality, reduced neural readiness/focus, underfueling, low hydration, soreness ≥3/5, or localized discomfort.",
        "Soreness 5/5, severe/acute pain, hesitation, or a worse warm-up sets maximalSprintRestricted=true.",
        "Yellow without a reassessment asks for better/same/worse after warm-up. Daily readiness does not rebuild the saved week.",
    ])
    add_callout(doc, "Important scope boundary", "The daily readiness check changes launch guidance. The weekly planner’s tier regression currently reads profile-level currentPain/return-to-training/taper only, and only in the Football/General branch.", "warning")

    add_schema_section(doc)

    add_heading(doc, "Complete deterministic weekly-plan generation", 1, "07  /  SELECTION ENGINE")
    add_numbered(doc, [
        "Respect plan mode. Logging-only and coach-plan modes return coach-managed output and do not generate a week.",
        "Derive season phase from the calendar or explicit override. A needs-calendar result stops generation rather than guessing.",
        "Resolve open weekdays. Preferred rest, practices, another sport, current-week competitions, and priority meets are blocked before matching.",
        "Choose the pathway branch: Track for primary track sport; Football for primary Football; General Speed for all other non-track sports.",
        "Build weekly roles. Track uses phase/event/experience architecture; Football/General uses tiered acceleration, upright speed, low support, and optional integration.",
        "Hard-gate candidates: Approved and validation-clean; matching pathway/event/level/phase; meet-window allowed; no specialized assisted/force-plate requirement.",
        "Rank deterministically. Category fit (40), event (25), phase (20), speed-goal overlap (10), duration fit (5), minus progressionLevel; preferred IDs receive priority; ID is the final tie-break.",
        "Pair strength on authored high-output roles, normally STR-01 then STR-02, consolidating hard work on quality days.",
        "Build PlannedWorkout snapshots and a Monday–Sunday schedule. Unselected days remain explained open/rest days.",
        "Return explanations, excluded harder options, setup, stop rule, and up to two approved alternatives. The current plan changes only after explicit save.",
    ])
    add_heading(doc, "Track general-preparation defaults", 2)
    add_table(doc, ["Role", "Load", "Default content"], [
        ["1", "High", "Acceleration + paired strength"],
        ["2", "Low", "Extensive tempo / capacity"],
        ["3", "High", "Maximum velocity + paired strength"],
        ["4", "Low", "Technical low day"],
        ["5", "High or moderate", "Speed endurance; beginner receives elastic power instead"],
    ], [0.6, 0.85, 5.25], 8)
    add_heading(doc, "Football and General Speed defaults", 2)
    add_table(doc, ["Days", "Roles"], [
        ["2", "Acceleration; upright speed"],
        ["3", "Acceleration; low technical/support; upright speed"],
        ["4", "Acceleration; low technical/support; upright speed; movement/trunk support"],
        ["5 beginner/developing", "Same first four plus elastic movement foundation"],
        ["5 trained/advanced", "Same first four plus 40-yard transfer or acceleration integration"],
        ["In season", "At most three roles: speed microdose, recovery support, reduced support"],
    ], [1.65, 5.05], 8)
    add_callout(doc, "Determinism", "Candidate scores, preferred-ID order, progression level, and record ID provide stable tie-breaking. No random dialogue or workout generation participates in plan selection.", "success")

    add_heading(doc, "Execution, tracking, completion, and logging", 1, "08  /  FROM RECORD TO LOG")
    add_table(doc, ["Library structure", "Current execution structure", "Saved behavior"], [
        ["Strength 3×3, 2×10, etc.", "strength tracking with sets and targetReps.", "One StrengthSetResult per set; load and reps can be recorded; quick completion snapshots preserve statuses."],
        ["Sprint 2×5 at 100 m", "track tracking with reps=10 and distance=100.", "Ten TrackRepResult items; set boundary and between-set rest are not represented."],
        ["Per-side work", "Usually text in notes/detail.", "No dedicated side key or left/right result array."],
        ["Timed hold", "If in strength section, duration often becomes target text/detail; otherwise completion-only.", "No generic per-set duration result field."],
        ["Distance repetition", "track tracking with rep count, distance, target intensity, and one restSeconds value.", "Per-rep completion/time/feeling can be saved."],
        ["Warm-up, plyometric, cooldown, many circuits", "completion tracking.", "Completed/skipped/pending only; authored dose remains in snapshot detail."],
        ["Whole-exercise quick completion", "Marks all child reps/sets completed and saves prior statuses.", "Undo restores statuses without deleting manual times, loads, notes, or feedback."],
    ], [1.45, 2.55, 2.7], 7.1)
    add_heading(doc, "Completion pipeline", 2)
    add_numbered(doc, [
        "Selected library records are converted to immutable PlannedWorkout snapshots.",
        "ActiveWorkoutSession initializes ActualExerciseResult entries and child track reps or strength sets.",
        "Rep/set/exercise controls update pending/completed/skipped state. Added or changed work retains an origin and change reason.",
        "Post-workout review saves completion, RPE, energy, soreness fields, optional notes, and a structured domain log when available.",
        "Completed history retains the planned snapshot, actual results, timestamps, readiness snapshot, and review.",
    ])

    add_heading(doc, "Exercise alternatives", 1, "09  /  RESOLUTION")
    doc.add_paragraph(
        "The weekly selector chooses whole workout alternatives deterministically. Inside several Approved records, exercise names still contain "
        "human-readable options such as “Back squat or trap-bar deadlift” or “Box jump or broad jump.” The current data model does not represent "
        "these as typed choice objects, and the planner does not resolve them from equipment."
    )
    add_table(doc, ["Alternative type", "Current behavior", "V3 status"], [
        ["Whole workout", "Planner ranks candidates and exposes up to two Approved alternatives; athlete can replace the day.", "Implemented"],
        ["Equipment substitution in item name", "Displayed as one executable label; athlete/coach interprets the option.", "Implemented but unresolved"],
        ["Typed exercise choice", "No alternative IDs/options array in LibraryWorkoutItem.", "Not implemented"],
        ["Per-side variant", "Often described in notes only.", "Not implemented as tracking structure"],
    ], [1.55, 3.25, 1.9], 7.4)
    add_callout(doc, "Conflict with requested ideal", "V3 cannot claim that every exercise alternative is deterministically resolved or selected through a structured UI. Several Approved records retain unresolved “or” labels. This is a validation and execution limitation.", "error")

    add_heading(doc, "Library approval validation", 1, "10  /  APPROVAL GATE")
    add_heading(doc, "Implemented checks", 2)
    add_bullets(doc, [
        "No “TBD” content anywhere in the serialized record.",
        "Name, purpose, and intended athlete are present.",
        "At least one pathway, event, athlete level, and season phase.",
        "At least one required surface.",
        "Workout-level intensity summary, recovery summary, coaching cues, safety notes, and source notes.",
        "All six workout sections exist and contain item arrays.",
        "Stored total and high-intensity sprint volumes exactly match calculated section-item volumes.",
        "Special-endurance records include 200m/400m event and phase restrictions.",
        "Advanced drop-jump names require an override and are rejected by the prototype.",
        "Overspeed/downhill names are rejected from Approved status.",
        "Approved workouts cannot be silently edited; revisions begin as Draft. Archived records restore to Draft.",
    ])
    add_heading(doc, "Requested rules not yet enforced", 2)
    add_bullets(doc, [
        "Unnamed or semantically vague circuits, mobility, warm-up, cooldown, or prehab blocks.",
        "“Quality reps” without a numeric dose.",
        "Undefined exercise references or typed exercise-registry lookups.",
        "Unresolved within-item alternatives such as “A or B.”",
        "Unsupported per-side, interval, block, or circuit execution structures.",
        "Item-level missing rest/intensity/completion metadata.",
        "Contradictory sets/reps versus execution tracking beyond sprint-volume arithmetic.",
    ])
    add_callout(doc, "Approval result", "All 64 current Approved records pass the implemented validator. Passing does not imply they satisfy the stronger planned rules above.", "warning")

    add_inventory(doc)
    add_regression_matrix(doc)
    add_sources(doc)

    add_heading(doc, "Known limitations and planned behavior", 1, "14  /  HONEST BOUNDARIES")
    add_table(doc, ["Area", "Implemented now", "Limitation / planned behavior"], [
        ["Track set fidelity", "Library keeps sets and reps.", "Execution flattens track sets; must add set-aware track tracking before claiming 2×5 fidelity."],
        ["Per-side work", "Text may state per side.", "No first-class side structure or side-specific result storage."],
        ["Readiness regression", "Football/General tier can regress and explains why.", "Track weekly tier does not regress; daily readiness only changes launch guidance."],
        ["Advanced track tiers", "Advanced eligibility is honored.", "Advanced 200/400 currently selects the same IDs as trained; advanced 60/100 mixes a GEN record."],
        ["200m pathway", "200m maps to long-sprint pathway.", "Weekly architecture’s isLongSprint check recognizes only 400m, so 200m uses short-style role preferences."],
        ["In-season support", "Practices/games are blocked and microdose week is generated.", "Tested non-track in-season case contains no meaningful loaded-strength exposure."],
        ["Exercise alternatives", "Readable alternatives appear in item names.", "No typed choice/resolution model."],
        ["Approval rigor", "Current structural and volume checks pass.", "Semantic dose/reference/execution checks requested for V3 are not implemented."],
        ["Sports", "Football has distinct pathway; others receive General Speed.", "No dedicated soccer/basketball/baseball/softball progressions."],
        ["Type safety", "Planner scripts pass.", "Repository typecheck has three app/profile.tsx errors because SprintEvent includes 300m while a helper accepts only 60/100/200/400."],
        ["Lint", "No lint defects were proven by this audit.", "expo lint did not return output within the audit window and was stopped; result is inconclusive."],
    ], [1.25, 2.55, 2.95], 6.8)
    add_callout(doc, "Release interpretation", "Planner regression scripts pass. V3 is complete as an implementation specification, but it deliberately does not certify the current working tree as free of all application defects.", "info")

    add_heading(doc, "Appendix A — Nonselectable records", 1, "A  /  DRAFT & ARCHIVED")
    rows = []
    for w in DATA["inventory"]:
        if w["approvalStatus"] != "approved":
            rows.append([
                w["id"],
                w["name"],
                w["approvalStatus"],
                w["primaryCategory"],
                "; ".join(w["approvalErrors"]) or "Not recommendation-eligible because status is not Approved.",
            ])
    add_table(doc, ["ID", "Name", "Status", "Category", "Reason / state"], rows, [0.75, 2.2, 0.75, 1.15, 1.9], 7)

    path = OUTPUT / "SprintLab_Workout_Library_and_Season_Engine_V3.docx"
    doc.save(path)
    return path


def build_audit() -> Path:
    doc = setup_document("SprintLab V3 Implementation Audit Report")
    add_cover(doc, "A concise evidence report for the repository, test runs, conflicts, and decisions behind V3.")

    add_heading(doc, "Audit conclusion", 1, "EXECUTIVE SUMMARY")
    doc.add_paragraph(
        "The deterministic planner and performance-time verification scripts pass in the reviewed working tree. "
        "The library contains 64 Approved recommendation-eligible records and generated 28 auditable profile runs. "
        "V3 documents the implementation faithfully and records material gaps instead of rewriting the app to match V2."
    )
    add_table(doc, ["Check", "Result", "Evidence"], [
        ["Planner architecture", "PASS", "Planner architecture scenarios passed."],
        ["Planner pathways/regressions", "PASS", "All release planner scenarios passed."],
        ["Performance-time ranges", "PASS", "Performance-time input range verification passed."],
        ["Library approval", "PASS", "64 Approved eligible; 0 Approved records with current validation errors."],
        ["TypeScript", "FAIL", "Three profile.tsx type errors involving SprintEvent=300m and a narrower helper."],
        ["Lint", "INCONCLUSIVE", "expo lint produced no result in the audit window and was stopped."],
        ["Profile matrix", "28 generated", "Beginner, trained, advanced, reduced-readiness, competition, and in-season cases."],
    ], [1.7, 1.1, 3.95], 7.4)

    add_heading(doc, "Repository state reviewed", 1)
    add_table(doc, ["Property", "Value"], [
        ["Repository", str(ROOT)],
        ["App version", "1.0.0"],
        ["Base commit", BASE_COMMIT],
        ["Base commit date", "July 26, 2026"],
        ["Review date", CREATED],
        ["Working tree", "Dirty; reviewed current filesystem state, not only the base commit."],
        ["Generated evidence", "docs/v3/build/v3-audit-data.json and test-logs/"],
    ], [1.6, 5.15], 8)

    add_heading(doc, "Files and modules inspected", 1)
    files = [
        "types/domain.ts — AthleteProfile, events, season calendar, performance/profile types",
        "types/workout-library.ts — exact library schema and approval/filter types",
        "types/index.ts — PlannedWorkout, execution tracking, active/completed session structures",
        "data/workout-library.ts — 70 seeded workout records and prescriptions",
        "data/workout-sources.ts — saved research-source catalog",
        "utils/plan-selector.ts — day blocking, hard gates, ranking, conversion, strength pairing, planner branches",
        "utils/weekly-architecture.ts — phase-specific high/low role architecture",
        "utils/season-engine.ts — season derivation and meet-window eligibility",
        "utils/workout-library.ts — persistence, migration, approval validation, archive/restore",
        "utils/readiness.ts — green/yellow/red readiness and maximal-sprint restriction",
        "utils/workout-session.ts — execution-item initialization and state derivation",
        "utils/domain-adapters.ts — completed-session/domain-log conversion",
        "app/workout.tsx — active execution, per-rep/set completion, quick exercise completion",
        "scripts/verify-planner-architecture.ts — architecture and strength-template regressions",
        "scripts/verify-plan-pathways.ts — pathway, tier, schedule-load, coach/log, and History regressions",
        "scripts/verify-performance-time.ts — performance-time range regressions",
        "SprintLab_Workout_Library_and_Season_Engine_V2.md — background only",
        "SprintLab_40_Yard_and_General_Speed_Research.md — scientific rationale and S-source catalog",
    ]
    add_bullets(doc, files, compact=True)

    add_heading(doc, "Tests run and results", 1)
    add_table(doc, ["Command / evidence", "Result", "Notes"], [
        ["npm run verify:planner", "PASS", "Architecture and all release pathway scenarios passed."],
        ["npm run verify:performance-time", "PASS", "Event-specific input range verification passed."],
        ["docs/v3/build/extract-v3-audit.ts", "PASS", "Generated 28 real selector runs and serialized all 70 library records."],
        ["npx tsc --noEmit", "FAIL", "Three pre-existing profile.tsx errors remain after excluding audit-script errors."],
        ["npm run lint", "INCONCLUSIVE", "No output/termination after extended wait; manually stopped."],
    ], [2.1, 1.05, 3.6], 7.6)

    add_heading(doc, "Implementation/documentation conflicts", 1)
    conflicts = [
        ["Sprint set structure", "Requested/V2-style contract expects 2×5 to remain two sets.", "Current plan conversion flattens to 10 track reps.", "Documented as unresolved; app not changed."],
        ["Per-side tracking", "Requested contract expects per-side structure.", "No typed side field or side result model.", "Documented as unresolved."],
        ["Readiness tier regression", "Requested rule applies across pathways.", "Only Football/General branch lowers a tier.", "Verified with reduced-readiness matrix; track cases fail."],
        ["Advanced track differentiation", "Advanced should express a distinct higher tier.", "Advanced 200/400 selects same IDs as trained; short path mixes GEN record.", "Marked Questionable."],
        ["200m architecture", "200/400 intended as one long-sprint pathway.", "isLongSprint(event) is true only for 400m.", "Documented as an implementation limitation."],
        ["Approval validation", "Requested semantic rules reject vague or unresolved items.", "Current validator is primarily structural/volume-based.", "V3 separates implemented checks from planned rules."],
        ["Equipment resolution", "Ideal behavior resolves equipment-specific alternatives.", "Ordinary access is assumed; several item names contain unresolved “or” options.", "Documented; no false claim of deterministic exercise resolution."],
        ["Source IDs", "Scientific sources should be externally verifiable.", "Two records use “User helper note.”", "Identified as an internal authoring note, not a publication."],
    ]
    add_table(doc, ["Topic", "Documentation/requirement", "Current code", "Decision"], conflicts, [1.2, 2.0, 2.05, 1.5], 6.6)

    add_heading(doc, "Profile-run decisions", 1)
    ratings = Counter(scenario_rating(s)[0] for s in DATA["scenarios"])
    doc.add_paragraph(
        f"Ratings across 28 runs: {ratings['PASS']} Pass, {ratings['QUESTIONABLE']} Questionable, {ratings['FAIL']} Fail. "
        "Ratings judge structural behavior, not one exact exercise choice."
    )
    rows = []
    for s in DATA["scenarios"]:
        rating, reason = scenario_rating(s)
        rows.append([
            s["scenario"]["name"],
            ", ".join(scenario_selected_ids(s)),
            rating,
            reason,
        ])
    add_table(doc, ["Profile", "Selected IDs", "Result", "Reason"], rows, [1.65, 2.7, 0.75, 1.65], 6.5)

    add_heading(doc, "Decisions made", 1)
    add_bullets(doc, [
        "Used the current repository and test output as implementation authority; V2 was not copied forward.",
        "Did not modify planner or application behavior to make it agree with the requested ideal.",
        "Documented broader AthleteProfile fields only when current selection code reads them.",
        "Separated daily readiness launch guidance from weekly profile-level tier regression.",
        "Listed all 64 Approved/selectable records with current prescriptions and all six nonselectable records separately.",
        "Used saved R and S sources for rationale without implying endorsement.",
        "Marked actual defects and incomplete semantics as limitations even when planner scripts passed.",
    ])

    add_heading(doc, "Could not be verified", 1)
    add_bullets(doc, [
        "A clean full TypeScript build, because three profile.tsx type errors remain.",
        "A completed lint run, because the configured lint process did not return a result.",
        "Physical-device rendering and interaction for every workout prescription; this audit is implementation/test based.",
        "Scientific support for the internal “User helper note” source label.",
        "Set-preserving sprint execution or per-side logging, because those structures are not implemented.",
    ])
    add_callout(doc, "Final audit judgment", "V3 is a faithful implementation specification. It is not a claim that every requested planner/execution behavior is complete.", "success")

    path = OUTPUT / "SprintLab_V3_Implementation_Audit_Report.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    OUTPUT.mkdir(parents=True, exist_ok=True)
    v3 = build_v3()
    audit = build_audit()
    print(v3)
    print(audit)
